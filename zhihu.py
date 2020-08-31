import requests
import re
from bs4 import BeautifulSoup
import os
import urllib
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Pt
import time
import json

#定义常量
homePath = "D:\\Work\\Document\\BaiduZhihu\\"
docPath = homePath+"发送缓存区\\"
imgCache = homePath+"cache\\imgcache.jpg"
logPath = homePath+"cache\\log.txt"
log2Path =homePath+"发送缓存区\\[log].txt"

#获取Html
def getHtmlText(url):
    try:
        headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.138 Safari/537.36'}
        r=requests.get(url,headers=headers,timeout=30)
        r.raise_for_status()
        r.encoding=r.apparent_encoding
        return r.text
    except:
        return "no"

#获取文件
def getDoc(URL):
    html=getHtmlText(URL)
    soup = BeautifulSoup(html,'lxml')
    soup = soup.article
    #获取标题和作者
    title = str(soup.find("h1","Post-Title").get_text())
    authorName = str(soup.find_all(attrs={"data-za-detail-view-element_name": "User"})[1].get_text())
    #添加log2文件
    with open(log2Path,"a",encoding="utf-8") as log2File:
        log2File.write("\n------------------------------\n")
        log2File.write("标题："+title+"\n链接："+URL+"\n")
        log2File.close

    logFile.write("----"+title+"----\n")
    #初始化python-docx 设置字体样式
    doc = Document()
    doc.styles['Normal'].font.name = u'等线'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'等线')

    #添加标题 作者
    doc_tittle=doc.add_heading('',level=1).add_run(title)
    doc_tittle.font.name = u'等线'
    doc_tittle._element.rPr.rFonts.set(qn('w:eastAsia'), u'等线')
    doc_tittle.font.size = Pt(24)
    doc_tittle.font.color.rgb=RGBColor(0,0,0)

    doc.add_paragraph("作者： "+authorName+" 来源：人工智能学习圈")
    doc.add_paragraph("本文未经授权，严禁转载，转载请联系作者本人。")
    doc.add_paragraph("")
    if(authorName=="逆暗" or authorName== "裴丘" or authorName== "hooo"): title="【#】"+title
    logFile.write("标题："+title+"\n链接："+URL+"\n")

    #缩小文字区
    maintext = soup.find(attrs={"class":"RichText ztext Post-RichText"})
    #遍历maintext所有tag
    for child in maintext.children:
        #标题或者正文，直接添加
        if(child.name == 'p' or child.name =='h2' or child.name =='h3' or child.name =='hr' or  child.name =='blockquote'):
            doc.add_paragraph(child.get_text())
        #代码块，后续手动截图
        elif(child.name =='div'):
            print("【CODE!!!】")
            doc.add_paragraph("【CODE！！】")
            logFile.write("【CODE】\n")
        #图片的操作
        elif(child.name =='figure'):
            img = child.find_all('img')[0]
            imgType=""
            if("data-original" in img.attrs): #原始图片
                imgType="data-original"
            elif("data-actualsrc" in img.attrs): #真实图片
                imgType="data-actualsrc"
            elif("src" in img.attrs): #只有路径
                imgType="src"
            else:
                imgType=""
                print("image ERROR!")
                print(img)
                doc.add_paragraph("IMG_ERROR!!!】")
                logFile.write("【IMG_ERROR】\n")
            if(imgType!=""):
                #print("imgType: "+imgType)
                #下载图片 添加图片到doc
                urllib.request.urlretrieve(str(img[imgType]),filename=imgCache)
                if(getFileSize(imgCache)<0.8): #图片太小跳过
                    print("imgsize:"+str(getFileSize(imgCache))+"kb")
                    logFile.write("imgsize:"+str(getFileSize(imgCache))+"kb\n")
                    urllib.request.urlretrieve("https://i0.hdslb.com/bfs/archive/487d59e34dc3dc5080b64e9337119b2dae050e45.png",filename=imgCache)
                doc.add_picture(imgCache)
                if(child.get_text()!=""):
                    print("Photo Annotate Added")
                    logFile.write("PhotoAnnotateAdded\n")
                    doc.add_paragraph("【"+child.get_text()+"】")
        #列表
        elif(child.name =='ul'):
            for li in child.find_all('li'):
                doc.add_paragraph(li.get_text())
        elif(child.name =='ol'):
            i=1
            for li in child.find_all('li'):
                doc.add_paragraph(str(i) +". "+ li.get_text())
                i=i+1
        #超链接
        elif(child.name =='a'):
             doc.add_paragraph(child['href'])
             if(child.get_text()!=""):
                    doc.add_paragraph(child.get_text())
        else:
            print("ERROR! something unrecord")
            logFile.write("【unrecordERROR!】\n")
    doc.add_paragraph("")
    doc.add_paragraph("原文链接:"+URL)
    
    #防止文件的名字出现不能保存的字符，替换一下。
    title=title.replace('|','')
    title=title.replace(':','')
    title=title.replace('/','')
    doc.save(docPath +title+".docx" )
    print(title+"  SAVED")
    print()
    logFile.write("\n\n")




#获取图片的大小
def getFileSize(filePath):
    fsize = os.path.getsize(filePath)
    fsize = fsize/float(1024)
    #print(round(fsize,2))
    return round(fsize,2)


#获取XueAI的URL
def getXueAI(zhuanlanAPIurl):
    html = getHtmlText(zhuanlanAPIurl)
    #直接导入json
    xueAIJsonData=json.loads(html) 
    
    #打开url缓存json文件
    with open(homePath+"cache\\urlListCache.json","r",encoding ="utf8") as urlfp:
        urlListCache=json.load(urlfp)       
        newUrl =[]                          #初始化新url
        for data in xueAIJsonData['data']:  #遍历专栏json文件，剔除旧的url
            xueaiUrl=data['url']
            for url in urlListCache['url']:
                if(xueaiUrl==url):
                    xueaiUrl=""
                    break
            if(xueaiUrl!=""):
                newUrl.append(xueaiUrl)
                print("检测到新URL: "+xueaiUrl)

        #将更新后的数据写入urlListCache.json
        with open(homePath+"cache\\urlListCache.json","w",encoding ="utf8") as urlfp:
            for i in newUrl:
                urlListCache['url'].append(i)
            json.dump(urlListCache,urlfp)
    if(len(newUrl) ==0): 
        print("没有新的URL")
        return []
    elif(len(newUrl) ==10):
        newUrl.extend(getXueAI(xueAIJsonData['paging']['next']))
        return newUrl
    else:return newUrl
    

if __name__ =='__main__':

 
    #打开log文件
    logFile = open(logPath,"a",encoding="utf-8")
    logFile.write("\n")
    logFile.write("---"+time.asctime( time.localtime(time.time()) )+"---\n")
    logFile.close
    log2File = open(log2Path,"a",encoding="utf-8")
    log2File.write("\n")
    log2File.write("---"+time.asctime( time.localtime(time.time()) )+"---\n")
    log2File.close

    urlList = getXueAI("https://zhuanlan.zhihu.com/api/columns/xueAI/articles?include=data%5B*%5D.admin_closed_comment%2Ccomment_count%2Csuggest_edit%2Cis_title_image_full_screen%2Ccan_comment%2Cupvoted_followees%2Ccan_open_tipjar%2Ccan_tip%2Cvoteup_count%2Cvoting%2Ctopics%2Creview_info%2Cauthor.is_following%2Cis_labeled%2Clabel_info")
    print("共"+ str(len(urlList))+"条")
    for url in urlList:
        if url != "":
            logFile = open(logPath,"a",encoding="utf-8")
            getDoc(url)
            logFile.close

    print("end")

